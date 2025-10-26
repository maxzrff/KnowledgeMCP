"""
DOCX document processor.
"""

from pathlib import Path
from typing import Any

import docx

from src.models.document import DocumentFormat
from src.processors.base import BaseProcessor
from src.utils.logging_config import get_logger

logger = get_logger(__name__)


class DOCXProcessor(BaseProcessor):
    """DOCX document processor."""

    @property
    def supported_format(self) -> DocumentFormat:
        return DocumentFormat.DOCX

    async def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX."""
        try:
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraphs)
            logger.info(f"Extracted {len(text)} characters from DOCX: {file_path.name}")
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise

    async def extract_metadata(self, file_path: Path) -> dict[str, Any]:
        """Extract metadata from DOCX."""
        try:
            doc = docx.Document(file_path)
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "format": "docx",
            }

            # Core properties
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.subject:
                metadata["subject"] = doc.core_properties.subject

            return metadata
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
            return {"format": "docx"}
